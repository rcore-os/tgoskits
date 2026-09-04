#!/usr/bin/env python3
"""Generate Task 1 weekly report (stage 1 & 2) as Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = "/home/allen/codes/os/tgoskits/plans/task1-reports/2026-揭榜挂帅-技术攻关周报-20260724.docx"


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_label(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_body(doc, text, size=12):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, size=size)


def main():
    doc = Document()

    add_heading(doc, "2026 揭榜挂帅擂台赛技术攻关周报", size=18)
    add_body(doc, "2026.7.24")
    add_body(doc, "1、玄枢智能系统团队")

    add_label(doc, "本周进展（当前研发进度、已解决的技术问题等）")

    progress = """
任务一（AxVisor 实时性改造与验证）阶段一、阶段二已完成主体落地，形成可复现的混合分区基线与调度改造能力，概要如下。

一、阶段一：基线复现与测量框架（已完成）

1. 混合分区拓扑（QEMU aarch64，-smp 4）
   • pCPU0：AxVisor 宿主；pCPU1–2：Linux 2vCPU 客户机（AI/stress 域）；pCPU3：RT 实时控制域（ArceOS 1vCPU）。
   • 新增 VM 配置：linux-smp2.toml（2vCPU，phys_cpu_ids=[1,2]）、arceos-rt-smp1.toml（pCPU3 独占）、zephyr-rt-baseline.toml、rtthread-rt-baseline.toml。

2. 测量与冒烟体系
   • 裸机 RTOS 抖动测量：rt-latency 测试（1ms/10ms 周期，200 样本，输出 mean/P99/max），已接入 cargo xtask arceos test。
   • AxVisor Linux 2vCPU 冒烟：test-suit/axvisor/normal/qemu/linux-smp2/ PASS。
   • RT 基线：Zephyr / RT-Thread QEMU smoke 用例与一键脚本（run-zephyr-rt-baseline.sh、run-rtthread-rt-baseline.sh）。
   • 任务一脚本集：os/axvisor/scripts/task1/（setup、run-mixed、build-arceos-rt-guest 等）。

3. 已解决的技术问题
   • 完成双客户机（Linux + RT）在 AxVisor 下的分区启动验证。
   • 建立裸机 vs 虚拟化客户机的统一 RT_LATENCY 日志格式，便于后续对比报告自动化。

二、阶段二：调度改造与虚拟化时序优化（已完成主体）

1. vCPU 优先级与宿主抢占调度
   • VM 配置新增 vcpu_priorities（Linux nice：RT 域 -20，Linux 域 +10）。
   • AxVisor / AxVM 启用 sched-cfs；spawn_vcpu_task 创建 vCPU 后应用优先级。
   • 中断 queue_interrupt 路径增加 wake_task，提升 RT 域被 Linux 抢占后的唤醒响应。

2. 架构定时器与 GIC 优化
   • passthrough_timer：VM exit 不再保存/恢复 CNTV/CNTP；CNTKCTL 允许 EL0 访问 counter/timer，满足 RTOS 用户态 tick。
   • 修复 RT 客户机 passthrough GIC 下 GICD ioremap 失败：rt-latency feature 补启 ax-std/paging，解决 mem_iomap Unsupported。
   • 验证：build-arceos-rt-guest.sh 构建 + arceos-rt-latency 用例 RT_LATENCY_PASS。

3. 阶段二验证数据（idle 短测矩阵，摘录）
   • 裸机 idle 1ms P99 ≈ 309312 ns；guest idle 改造后 1ms P99 ≈ 262656 ns。
   • stress 长稳 pre/post 对比：1ms P99 改善约 1.9%，10ms P99 改善约 5.6%（仅 vcpu_priorities 维度，距赛题 ≥50% 目标仍有差距，已纳入阶段三专项）。

4. 自动化与报告
   • collect-rt-latency-report.sh、collect-task1-matrix-report.sh：裸机/guest 改造前后对比报告归档至 plans/task1-reports/。
   • mixed-rt-stress-round1 长稳用例（180k 样本）可一键复现。
"""
    add_body(doc, progress)

    add_label(doc, "下周计划")
    next_plan = """
1. 任务一阶段三：补齐 stress 压力下改造前后完整对比，探索 IRQ 虚拟注入延迟 benchmark 与额外调度/隔离优化，争取向赛题 ≥50% P99 改善目标推进。
2. 任务一交卷材料：整理设计说明、对比图表与 SUBMISSION-STATUS 终稿。
3. 任务二（客户机间通信）：在 icpc-smoke 通过后，继续可靠性（ACK/重传/心跳）与故障注入用例。
"""
    add_body(doc, next_plan)

    add_label(doc, "代码链接：")
    links = """
• 仓库：https://github.com/rcore-os/tgoskits（Task 1 相关分支/本地工作区）
• 实施记录：plans/task1-实施记录.md
• 任务一指南：os/axvisor/doc/task1-realtime.md
• 测试报告：plans/task1-reports/
• 关键配置：os/axvisor/configs/vms/qemu/aarch64/linux-smp2.toml、arceos-rt-smp1.toml
"""
    add_body(doc, links)

    add_label(doc, "需协助事项")
    assist = """
1. 如需在物理板卡（非 QEMU）上复现 Task 1 混合分区，需确认板端 SMP 核数与 pCPU 绑定策略是否可沿用当前四核 QEMU 拓扑。
2. 赛题「改造前基线」若要求完全未改动的 AxVisor 版本对比，需明确是否提供官方参考 tag 或允许以 pre-opt 配置（去掉 vcpu_priorities）作为改造前对照。
3. 暂无其他阻塞项；阶段三 stress ≥50% 改善需团队内部继续攻关，暂不依赖外部资源。
"""
    add_body(doc, assist)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
